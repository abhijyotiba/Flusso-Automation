"""
Attachment Processor Utility
Extracts text content from various attachment types (PDF, DOCX, XLSX, TXT, etc.)
Primary focus: PDF processing using PyMuPDF

NOTE: Freshdesk attachment URLs can be:
  1. Direct API URLs - require Freshdesk API key auth
  2. S3 signed URLs (amazonaws.com) - already contain AWS signature, NO auth needed
  
Adding auth to S3 signed URLs causes HTTP 400 errors!
"""

import logging
import io
import requests
from requests.auth import HTTPBasicAuth
import time
from typing import Dict, Any, List, Optional, Tuple
from dataclasses import dataclass

from app.config.settings import settings

logger = logging.getLogger(__name__)

# Lazy imports for optional dependencies
fitz = None  # PyMuPDF
Document = None  # python-docx
load_workbook = None  # openpyxl


def _import_pymupdf():
    global fitz
    if fitz is None:
        import fitz as _fitz
        fitz = _fitz
    return fitz


def _import_docx():
    global Document
    if Document is None:
        from docx import Document as _Document
        Document = _Document
    return Document


def _import_openpyxl():
    global load_workbook
    if load_workbook is None:
        from openpyxl import load_workbook as _load_workbook
        load_workbook = _load_workbook
    return load_workbook


@dataclass
class AttachmentContent:
    """Represents extracted content from an attachment"""
    filename: str
    file_type: str
    content: str
    page_count: int = 0
    extraction_time: float = 0.0
    error: Optional[str] = None
    size_bytes: int = 0


# Supported content types mapping
SUPPORTED_TYPES = {
    # PDFs (PRIMARY FOCUS)
    "application/pdf": "pdf",
    "application/x-pdf": "pdf",
    
    # Word documents
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "application/msword": "doc",
    
    # Excel spreadsheets
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": "xlsx",
    "application/vnd.ms-excel": "xls",
    
    # Text files
    "text/plain": "txt",
    "text/csv": "csv",
    "text/html": "html",
    
    # Images (just log, don't extract text)
    "image/jpeg": "image",
    "image/png": "image",
    "image/gif": "image",
    "image/webp": "image",
}


def download_attachment(url: str, timeout: int = 30) -> Tuple[Optional[bytes], Optional[str]]:
    """
    Download attachment from Freshdesk URL.
    
    IMPORTANT: Freshdesk uses two types of attachment URLs:
    1. Direct Freshdesk API URLs - require Basic Auth with API key
    2. S3 signed URLs (amazonaws.com) - already contain AWS signature, NO auth needed
    
    Adding auth to S3 signed URLs causes HTTP 400 errors!
    
    Returns:
        Tuple of (file_bytes, error_message)
    """
    try:
        logger.info(f"📥 Downloading attachment: {url[:100]}...")
        
        # Detect if this is an S3 signed URL (contains amazonaws.com and signature)
        # S3 URLs already have authentication built into the URL (X-Amz-Signature)
        is_s3_signed_url = (
            "amazonaws.com" in url.lower() or 
            "X-Amz-Signature" in url or
            "x-amz-signature" in url.lower()
        )
        
        if is_s3_signed_url:
            # S3 signed URLs should NOT use authentication - it causes 400 errors
            logger.info(f"📥 Detected S3 signed URL - downloading without auth")
            response = requests.get(
                url,
                timeout=timeout,
                stream=True
            )
        else:
            # Direct Freshdesk API URLs require Basic Auth
            logger.info(f"📥 Direct Freshdesk URL - using API key auth")
            auth = HTTPBasicAuth(settings.freshdesk_api_key, "X")
            response = requests.get(
                url, 
                auth=auth,
                timeout=timeout,
                stream=True
            )
        
        response.raise_for_status()
        
        # Get file size for logging
        file_size = len(response.content)
        logger.info(f"✅ Downloaded {file_size / 1024:.1f} KB")
        
        return response.content, None
    except requests.exceptions.Timeout:
        return None, f"Download timeout after {timeout}s"
    except requests.exceptions.HTTPError as e:
        if e.response.status_code == 401:
            return None, "Authentication failed - check Freshdesk API key"
        elif e.response.status_code == 403:
            return None, "Access forbidden - S3 URL may have expired"
        elif e.response.status_code == 404:
            return None, "Attachment not found (may have been deleted)"
        elif e.response.status_code == 400:
            return None, "Bad request - URL may be malformed or expired"
        return None, f"HTTP error: {e.response.status_code}"
    except requests.exceptions.RequestException as e:
        return None, f"Download failed: {str(e)}"


def extract_pdf_text(file_bytes: bytes, filename: str, max_pages: int = 50) -> AttachmentContent:
    """
    Extract text from PDF using PyMuPDF (fitz).
    
    This is the PRIMARY extraction method since 95% of attachments are PDFs.
    PyMuPDF is fast and handles most PDF types well.
    """
    start_time = time.time()
    
    try:
        fitz = _import_pymupdf()
        
        # Open PDF from bytes
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        page_count = len(doc)
        
        logger.info(f"📄 Processing PDF: {filename} ({page_count} pages)")
        
        text_parts = []
        pages_processed = min(page_count, max_pages)
        
        for page_num in range(pages_processed):
            page = doc[page_num]
            page_text = page.get_text()
            
            if page_text.strip():
                text_parts.append(f"--- Page {page_num + 1} ---\n{page_text}")
        
        doc.close()
        
        content = "\n\n".join(text_parts)
        extraction_time = time.time() - start_time
        
        if not content.strip():
            # PDF might be scanned/image-based
            return AttachmentContent(
                filename=filename,
                file_type="pdf",
                content="[PDF contains images/scanned content - no extractable text]",
                page_count=page_count,
                extraction_time=extraction_time,
                size_bytes=len(file_bytes),
                error="Image-based PDF"
            )
        
        logger.info(f"✅ Extracted {len(content)} chars from {pages_processed} pages in {extraction_time:.2f}s")
        
        return AttachmentContent(
            filename=filename,
            file_type="pdf",
            content=content,
            page_count=page_count,
            extraction_time=extraction_time,
            size_bytes=len(file_bytes)
        )
        
    except Exception as e:
        logger.error(f"❌ PDF extraction failed for {filename}: {e}", exc_info=True)
        return AttachmentContent(
            filename=filename,
            file_type="pdf",
            content="",
            extraction_time=time.time() - start_time,
            size_bytes=len(file_bytes),
            error=str(e)
        )


def extract_docx_text(file_bytes: bytes, filename: str) -> AttachmentContent:
    """Extract text from Word documents (.docx)"""
    start_time = time.time()
    
    try:
        Document = _import_docx()
        
        doc = Document(io.BytesIO(file_bytes))
        
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        
        content = "\n\n".join(paragraphs)
        extraction_time = time.time() - start_time
        
        logger.info(f"✅ Extracted {len(content)} chars from DOCX in {extraction_time:.2f}s")
        
        return AttachmentContent(
            filename=filename,
            file_type="docx",
            content=content,
            extraction_time=extraction_time,
            size_bytes=len(file_bytes)
        )
        
    except Exception as e:
        logger.error(f"❌ DOCX extraction failed for {filename}: {e}", exc_info=True)
        return AttachmentContent(
            filename=filename,
            file_type="docx",
            content="",
            extraction_time=time.time() - start_time,
            size_bytes=len(file_bytes),
            error=str(e)
        )


def extract_xlsx_text(file_bytes: bytes, filename: str, max_rows: int = 500) -> AttachmentContent:
    """Extract text from Excel spreadsheets (.xlsx)"""
    start_time = time.time()
    
    try:
        load_workbook = _import_openpyxl()
        
        wb = load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
        
        sheets_content = []
        
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            rows_text = []
            row_count = 0
            
            for row in sheet.iter_rows(values_only=True):
                if row_count >= max_rows:
                    rows_text.append(f"... (truncated at {max_rows} rows)")
                    break
                    
                # Filter out None values and convert to strings
                row_values = [str(cell) for cell in row if cell is not None]
                if row_values:
                    rows_text.append(" | ".join(row_values))
                    row_count += 1
            
            if rows_text:
                sheets_content.append(f"=== Sheet: {sheet_name} ===\n" + "\n".join(rows_text))
        
        wb.close()
        
        content = "\n\n".join(sheets_content)
        extraction_time = time.time() - start_time
        
        logger.info(f"✅ Extracted {len(content)} chars from XLSX in {extraction_time:.2f}s")
        
        return AttachmentContent(
            filename=filename,
            file_type="xlsx",
            content=content,
            extraction_time=extraction_time,
            size_bytes=len(file_bytes)
        )
        
    except Exception as e:
        logger.error(f"❌ XLSX extraction failed for {filename}: {e}", exc_info=True)
        return AttachmentContent(
            filename=filename,
            file_type="xlsx",
            content="",
            extraction_time=time.time() - start_time,
            size_bytes=len(file_bytes),
            error=str(e)
        )


def extract_text_file(file_bytes: bytes, filename: str) -> AttachmentContent:
    """Extract content from plain text files"""
    start_time = time.time()
    
    try:
        # Try different encodings
        for encoding in ['utf-8', 'latin-1', 'cp1252']:
            try:
                content = file_bytes.decode(encoding)
                break
            except UnicodeDecodeError:
                continue
        else:
            content = file_bytes.decode('utf-8', errors='replace')
        
        extraction_time = time.time() - start_time
        
        logger.info(f"✅ Extracted {len(content)} chars from text file in {extraction_time:.2f}s")
        
        return AttachmentContent(
            filename=filename,
            file_type="txt",
            content=content,
            extraction_time=extraction_time,
            size_bytes=len(file_bytes)
        )
        
    except Exception as e:
        logger.error(f"❌ Text extraction failed for {filename}: {e}", exc_info=True)
        return AttachmentContent(
            filename=filename,
            file_type="txt",
            content="",
            extraction_time=time.time() - start_time,
            size_bytes=len(file_bytes),
            error=str(e)
        )


def process_attachment(attachment: Dict[str, Any]) -> Optional[AttachmentContent]:
    """
    Process a single attachment and extract its text content.
    
    Args:
        attachment: Freshdesk attachment dict with keys like 'attachment_url', 'content_type', 'name'
        
    Returns:
        AttachmentContent with extracted text, or None if unsupported/failed
    """
    url = attachment.get("attachment_url") or attachment.get("url")
    content_type = str(attachment.get("content_type", "")).lower()
    filename = attachment.get("name", "unknown")
    
    if not url:
        logger.warning(f"⚠️ Attachment {filename} has no URL")
        return None
    
    # Determine file type
    file_type = SUPPORTED_TYPES.get(content_type)
    
    # Fallback: check file extension
    if not file_type:
        ext = filename.lower().split('.')[-1] if '.' in filename else ''
        ext_map = {
            'pdf': 'pdf',
            'docx': 'docx',
            'doc': 'doc',
            'xlsx': 'xlsx',
            'xls': 'xls',
            'txt': 'txt',
            'csv': 'csv',
            'html': 'html',
            'jpg': 'image',
            'jpeg': 'image',
            'png': 'image',
            'gif': 'image',
        }
        file_type = ext_map.get(ext)
    
    if not file_type:
        logger.info(f"⏭️ Skipping unsupported attachment type: {content_type} ({filename})")
        return None
    
    # Skip images (handled separately by vision pipeline)
    if file_type == "image":
        logger.debug(f"⏭️ Skipping image attachment (handled by vision pipeline): {filename}")
        return None
    
    # Download the file
    file_bytes, error = download_attachment(url)
    if error:
        logger.error(f"❌ Failed to download {filename}: {error}")
        return AttachmentContent(
            filename=filename,
            file_type=file_type,
            content="",
            error=error
        )
    
    # Extract text based on type
    if file_type == "pdf":
        return extract_pdf_text(file_bytes, filename)
    elif file_type == "docx":
        return extract_docx_text(file_bytes, filename)
    elif file_type in ("xlsx", "xls"):
        return extract_xlsx_text(file_bytes, filename)
    elif file_type in ("txt", "csv", "html"):
        return extract_text_file(file_bytes, filename)
    elif file_type == "doc":
        # Old .doc format - limited support
        logger.warning(f"⚠️ Old .doc format has limited support: {filename}")
        return AttachmentContent(
            filename=filename,
            file_type="doc",
            content="[Old .doc format - please convert to .docx for text extraction]",
            error="Legacy format"
        )
    
    return None


def process_all_attachments(attachments: List[Dict[str, Any]]) -> Dict[str, Any]:
    """
    Process all attachments from a ticket and extract text content.
    
    Args:
        attachments: List of Freshdesk attachment dicts
        
    Returns:
        Dict with:
            - extracted_content: Combined text from all attachments
            - attachment_summary: List of processed attachment info
            - images: List of image URLs (for vision pipeline)
            - stats: Processing statistics
    """
    if not attachments:
        return {
            "extracted_content": "",
            "attachment_summary": [],
            "images": [],
            "stats": {"total": 0, "processed": 0, "failed": 0, "images": 0}
        }
    
    logger.info(f"📎 Processing {len(attachments)} attachment(s)...")
    start_time = time.time()
    
    extracted_contents: List[AttachmentContent] = []
    images: List[str] = []
    failed_count = 0
    
    for att in attachments:
        if not isinstance(att, dict):
            continue
        
        content_type = str(att.get("content_type", "")).lower()
        url = att.get("attachment_url") or att.get("url")
        
        # Collect images for vision pipeline
        if content_type.startswith("image/"):
            if url:
                images.append(url)
            continue
        
        # Process document attachments
        result = process_attachment(att)
        if result:
            if result.error and not result.content:
                failed_count += 1
            extracted_contents.append(result)
    
    # Build combined content
    content_parts = []
    attachment_summary = []
    
    for idx, content in enumerate(extracted_contents, 1):
        summary = {
            "filename": content.filename,
            "type": content.file_type,
            "chars": len(content.content),
            "pages": content.page_count,
            "error": content.error
        }
        attachment_summary.append(summary)
        
        if content.content:
            header = f"\n{'='*60}\n📄 ATTACHMENT {idx}: {content.filename} ({content.file_type.upper()})"
            if content.page_count:
                header += f" - {content.page_count} pages"
            header += f"\n{'='*60}\n"
            content_parts.append(header + content.content)
    
    combined_content = "\n".join(content_parts)
    total_time = time.time() - start_time
    
    stats = {
        "total": len(attachments),
        "processed": len(extracted_contents),
        "failed": failed_count,
        "images": len(images),
        "total_chars": len(combined_content),
        "processing_time": total_time
    }
    
    logger.info(f"✅ Attachment processing complete: {stats['processed']} docs, {stats['images']} images, {stats['total_chars']} chars in {total_time:.2f}s")
    
    return {
        "extracted_content": combined_content,
        "attachment_summary": attachment_summary,
        "images": images,
        "stats": stats
    }
